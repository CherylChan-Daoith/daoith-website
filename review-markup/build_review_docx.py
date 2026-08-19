#!/usr/bin/env python3
"""Build review Word files: red = insert, red+strikethrough = delete (keep original)."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = Path("/home/ubuntu/.cursor/projects/workspace/uploads")
OUT = Path("/workspace/review-markup")
ART = Path("/opt/cursor/artifacts")

LEGEND = (
    "【审阅说明】红色字体＝建议新增；红色+删除线＝建议删除（原文仍保留，便于对照）。"
    "本文件仅供审阅，尚未写入知识库 / Dify。请看过后再决定是否采用。"
)


def rpr_of(run):
    if run is None:
        return None
    rPr = run._element.find(qn("w:rPr"))
    return deepcopy(rPr) if rPr is not None else None


def apply_kind(run, kind, base_rpr):
    r = run._element
    old = r.find(qn("w:rPr"))
    if old is not None:
        r.remove(old)
    if base_rpr is not None:
        r.insert(0, deepcopy(base_rpr))
    rPr = r.get_or_add_rPr()
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    for a in list(color.attrib):
        del color.attrib[a]
    if kind in ("ins", "del"):
        color.set(qn("w:val"), "FF0000")
    else:
        color.set(qn("w:val"), "141414")
    strike = rPr.find(qn("w:strike"))
    if kind == "del":
        if strike is None:
            strike = OxmlElement("w:strike")
            rPr.append(strike)
        strike.set(qn("w:val"), "true")
    elif strike is not None:
        rPr.remove(strike)


def apply_parts(paragraph, parts, template_run):
    base = rpr_of(template_run)
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    for kind, text in parts:
        if text == "":
            continue
        run = paragraph.add_run(text)
        apply_kind(run, kind, base)


def replace_para(p, parts, tmpl=None):
    apply_parts(p, parts, tmpl or (p.runs[0] if p.runs else None))


def insert_after(paragraph, paragraphs_parts, template_run=None):
    tmpl = template_run or (paragraph.runs[0] if paragraph.runs else None)
    ref = paragraph._p
    last = None
    pPr = paragraph._p.find(qn("w:pPr"))
    for parts in paragraphs_parts:
        new_p = OxmlElement("w:p")
        if pPr is not None:
            new_p.append(deepcopy(pPr))
        ref.addnext(new_p)
        ref = new_p
        para = Paragraph(new_p, paragraph._parent)
        apply_parts(para, parts, tmpl)
        last = para
    return last


def insert_before(paragraph, parts, template_run=None):
    tmpl = template_run or (paragraph.runs[0] if paragraph.runs else None)
    pPr = paragraph._p.find(qn("w:pPr"))
    new_p = OxmlElement("w:p")
    if pPr is not None:
        new_p.append(deepcopy(pPr))
    paragraph._p.addprevious(new_p)
    para = Paragraph(new_p, paragraph._parent)
    apply_parts(para, parts, tmpl)
    return para


def find_exact(doc, text):
    hits = [p for p in doc.paragraphs if p.text == text]
    if len(hits) != 1:
        raise KeyError(f"expected 1 match, got {len(hits)} for: {text[:80]!r}")
    return hits[0]


def find_startswith(doc, prefix):
    hits = [p for p in doc.paragraphs if p.text.startswith(prefix)]
    if len(hits) != 1:
        raise KeyError(f"expected 1 startswith, got {len(hits)} for: {prefix[:80]!r}")
    return hits[0]


def save(doc, name):
    OUT.mkdir(parents=True, exist_ok=True)
    ART.mkdir(parents=True, exist_ok=True)
    p1 = OUT / name
    p2 = ART / name
    doc.save(p1)
    doc.save(p2)
    print("wrote", p1, "and", p2)


def body_template_run(doc):
    for para in doc.paragraphs:
        if para.runs:
            return para.runs[0]
    return None


def add_legend(doc, tmpl):
    first = doc.paragraphs[0]
    insert_before(first, [("ins", LEGEND)], template_run=tmpl)
    insert_before(first, [("keep", "")], template_run=tmpl)


def build_prompt():
    doc = Document(SRC / "___0819_7c0f.docx")
    tmpl = body_template_run(doc)
    add_legend(doc, tmpl)

    p = find_exact(
        doc,
        "4. 追问阶段总回复 **2～4 行**；问诊阶段默认不检索；仅当第六步：了解产品类别用户答「其他」并给出具体品名时，允许查出口退税率知识库，查完只用来归类，不写报告。",
    )
    replace_para(
        p,
        [
            ("keep", "4. 追问阶段总回复 **2～4 行**；问诊阶段默认不检索；"),
            (
                "del",
                "仅当第六步：了解产品类别用户答「其他」并给出具体品名时，允许查出口退税率知识库，查完只用来归类，不写报告。",
            ),
            (
                "ins",
                "仅允许两类检索：①第六步用户答「其他」并给出具体品名时，可查出口退税率知识库，查完只用来归类、不写报告；②诊断过程中用户突然问无关政策问题时，可检索知识库作简短回答后立刻回到当前诊断步骤。禁止其它检索。",
            ),
        ],
        tmpl,
    )

    p = find_exact(
        doc,
        "11. 诊断过程中若用户突然问无关具体问题：先用知识库**简短**回答，再回到**当前诊断步骤**继续问。",
    )
    replace_para(
        p,
        [
            (
                "keep",
                "11. 诊断过程中若用户突然问无关具体问题：先用知识库**简短**回答，再回到**当前诊断步骤**继续问。",
            ),
            ("ins", "（此即铁律4的第②类检索例外，答完必须回到当前步，禁止借机跳步或提前出报告。）"),
        ],
        tmpl,
    )

    p = find_exact(doc, "12．步号写死：2=主体，3=发货，6=产品，7=销售额（模型会把步号写错）")
    replace_para(
        p,
        [
            ("del", "12．"),
            ("ins", "12. "),
            ("keep", "步号写死：2=主体，3=发货，6=产品，7=销售额（模型会把步号写错）"),
        ],
        tmpl,
    )

    p = find_exact(doc, "「2. 您平台店铺的注册主体是哪一种？（可在下方点选）」")
    insert_after(
        p,
        [[
            (
                "ins",
                "官网基础选项（仅供你采信用户点选，禁止写入问句）：中国大陆公司、中国个人、个体户、中国香港公司、外籍个人、海外本土公司、其他境外公司。",
            )
        ]],
        tmpl,
    )

    p = find_exact(doc, "「3. 请问您的发货方式是以下哪一种？（可在下方点选）」")
    insert_after(
        p,
        [
            [("ins", "根据第一步用户选择的平台，官网选项对应为（仅供采信，禁止写入问句）：")],
            [("ins", "1. **亚马逊**：亚马逊FBA、自发货（国内直发）、自发货（海外仓发货）")],
            [("ins", "2. **阿里国际站**：自营出口、一达通代理出口3+N、一达通代理出口2+N、市场采购出口、便捷发货出口")],
            [("ins", "3. **SHEIN**：供货 SHEIN（国内仓）、供货 SHEIN（保税仓）、SHEIN平台入驻商家（国内直发）、SHEIN平台入驻商家（海外仓发货）")],
            [("ins", "4. **速卖通**：全托管（国内仓）、半托管（国内仓）、半托管（海外仓）、POP（国内直发）、POP（海外仓发货）")],
            [("ins", "5. **Temu**：全托管（国内仓）、半托管（国内仓）、半托管（海外仓）、POP（国内直发）、POP（海外仓发货）")],
            [("ins", "6. **Shopee**：全托管（国内仓）、Shopee海外仓、自发货（国内直发）、自发货（海外仓发货）")],
            [("ins", "7. **Lazada**：全托管（国内仓）、FBL海外仓、自发货（国内直发）、自发货（海外仓发货）")],
            [("ins", "8. **其他平台**：发货到平台海外仓、发货到平台国内仓、自发货（国内直发）、自发货（海外仓发货）")],
        ],
        tmpl,
    )

    p = find_exact(doc, "「4. 您目前货物的出口方式是怎么样的？（可在下方点选）」")
    insert_after(
        p,
        [[
            (
                "ins",
                "官网选项为（仅供采信，禁止写入问句）：正式报关出口（0110/9710）、正式报关出口（9810）、小包快递出口（9610/1210）、小包快递出口（未报关）、市场采购出口（1039）、委托货代出口、由平台安排出口、其他。",
            )
        ]],
        tmpl,
    )

    p = find_exact(doc, "「6. 您的产品属于以下哪种类别？（可在下方点选）」")
    insert_after(
        p,
        [
            [("ins", "官网选项为（仅供采信，禁止写入问句）：")],
            [("ins", "- 普货，能正常报关出口和退税")],
            [("ins", "- 0退税率产品（如贵重金属、珠宝玉石、钢材、铝材、玻璃、木材）")],
            [("ins", "- 产品涉及商检（如食品、化妆品、危险化学品、木制品、医疗用品）")],
            [("ins", "- 产品涉及海关备案商标但暂未获得授权")],
            [("ins", "- 其他（不在以上分类）")],
        ],
        tmpl,
    )

    p = find_exact(doc, "「7. 您目前年销售额约多少人民币？（可在下方点选）」")
    insert_after(
        p,
        [[
            (
                "ins",
                "官网选项为（仅供采信，禁止写入问句）：500万以下、500-2000万、2000-5000万、5000万-1亿、1-4亿、4-10亿、10亿以上。",
            )
        ]],
        tmpl,
    )

    p = find_startswith(doc, "第 1～7 步齐全后**不要再提问**")
    replace_para(
        p,
        [
            (
                "keep",
                "第 1～7 步齐全后**不要再提问**，必须采信对话中已确认的平台/主体/发货/出口/发票/产品类别/销售额（含系统已自动记入的「由平台安排出口」），检索知识库并**综合推理**后",
            ),
            ("del", "按下列标题输出报告："),
            ("ins", "按《出报告硬约束与路径要点》规定的四章输出报告，本提示词不再复述路径细则、样本文件名与章节标题。"),
        ],
        tmpl,
    )

    p = find_exact(
        doc,
        "收集齐第 1～7 步答案后，按下列顺序检索与推理，勿跳步、勿颠倒，：",
    )
    replace_para(
        p,
        [
            ("keep", "收集齐第 1～7 步答案后，按下列顺序检索与推理，勿跳步、勿颠倒"),
            ("del", "，："),
            ("ins", "。检索顺序与版式细节以《出报告硬约束与路径要点》为准。"),
        ],
        tmpl,
    )

    p = find_startswith(doc, "判断路径：生成报告前必读")
    replace_para(
        p,
        [
            ("keep", "判断路径：生成报告前必读`出报告硬约束与路径要点`判定路径"),
            (
                "del",
                "：根据用户对平台/主体/发货/出口/发票/产品/销售额的具体选项，判断属于路径 A / B / C / D（禁止对用户输出路径名）",
            ),
            ("ins", "（禁止对用户输出路径名；判定规则以该文件为准，此处不展开）。"),
        ],
        tmpl,
    )

    p = find_startswith(doc, "参考对应方案样本：")
    replace_para(
        p,
        [
            (
                "del",
                p.text,
            ),
            (
                "ins",
                "参考对应方案样本：按《出报告硬约束与路径要点》指定的【方案样本】路径A/B/C/D改写；**按用户档案改写，禁止整篇照抄样本画像**，禁止对用户写出路径名或样本里的示例平台/销售额。",
            ),
        ],
        tmpl,
    )

    p = find_startswith(doc, "生成合规诊断报告：四固定章节")
    replace_para(
        p,
        [
            ("keep", "生成合规诊断报告：四固定章节；实操附件等仅按需补强；"),
            (
                "del",
                "报告章节标题**只能**用四个固定名：`【核心风险诊断】` `【合规方案】` `【行动建议】` `【注意事项】`。",
            ),
            (
                "ins",
                "四章标题与版式以《出报告硬约束与路径要点》为准，禁止自拟「第一部分」「核心风险提示」「合规方案建议」等变体。",
            ),
        ],
        tmpl,
    )

    p = find_exact(
        doc,
        "1. **【变化点】**：用项目符号逐条写「字段：旧值 → 新值」。用户本轮原话是权威来源。",
    )
    replace_para(
        p,
        [
            ("keep", "1. **【变化点】**：用项目符号逐条写「字段：旧值 → 新值」。用户本轮原话是权威来源。"),
            (
                "ins",
                "例如用户写「产品属于 13% 退税率产品」，即使旧档案是 0 退税率，也必须写成「产品退税率：0% → 13%（按普货可退税处理）」。",
            ),
        ],
        tmpl,
    )

    p = find_startswith(doc, "3. 再按**新事实**输出完整更新版")
    insert_after(
        p,
        [
            [("ins", "铁律：")],
            [("ins", "- 本轮明确的新事实 **覆盖** 旧档案，禁止把旧结论原样再贴一遍。")],
            [("ins", "- 用户已声明非 0 退税率（如 13%）时，禁止仍按 0 退税率写方案、禁止出现「叠加 0 退税率产品」等与新事实矛盾的句子。")],
            [("ins", "- 用户声明 0 退税率时，才走 0 退税率专项评估。")],
            [("ins", "- 禁止把增值税 13% 当成出口退税率；只有用户明确说「退税率 / 出口退税率」时才按退税率采信。")],
            [("ins", "- 禁止用旧档案里的 0 退税率覆盖用户本轮声明的非 0 退税率。")],
        ],
        tmpl,
    )

    p = find_exact(doc, "- 禁止再说「第 1～7 步已齐，请基于诊断档案出报告」而忽略用户本轮新问题。")
    insert_after(
        p,
        [[
            (
                "ins",
                "- 禁止用旧档案里的 0 退税率覆盖用户本轮声明的 13%（或其他非 0）退税率。",
            )
        ]],
        tmpl,
    )

    save(doc, "提示词0819-修订对照.docx")


def build_hard():
    doc = Document(SRC / "____________868e.docx")
    tmpl = body_template_run(doc)
    add_legend(doc, tmpl)

    p = find_exact(doc, "11. 禁止把「业务流程一行」「3-5 条风险」「Let me write」这类说明词写进正文")
    insert_after(
        p,
        [[
            (
                "ins",
                "12. 若检索结果不足或冲突，以本篇路径要点 + 内置合规架构速查 + 用户已确认选项为准，禁止因检索失败而空泛作答。",
            )
        ]],
        tmpl,
    )

    p = find_exact(
        doc,
        "9.「供货 SHEIN（国内仓）」禁止写成「保税仓」；销售额档位禁止改写",
    )
    replace_para(
        p,
        [
            ("keep", "9.「供货 SHEIN（国内仓）」禁止写成「保税仓」"),
            ("ins", "；「供货 SHEIN（保税仓）」禁止写成「国内仓」；国内仓与保税仓不得互相改写"),
            ("keep", "；销售额档位禁止改写"),
        ],
        tmpl,
    )

    p = find_exact(
        doc,
        "- 【合规方案】对照说明写成并列卡片（不要嵌套 a./b.）：",
    )
    replace_para(
        p,
        [
            ("keep", "- 【合规方案】对照说明写成并列卡片（不要嵌套 a./b.）"),
            ("ins", "；非定制／定制两条并列卡片仅路径B需要时使用，路径A/C/D不要套用该对照"),
            ("keep", "："),
        ],
        tmpl,
    )
    p = find_exact(doc, "    - **非定制类产品**：……")
    replace_para(
        p,
        [
            ("ins", "    （仅路径B）"),
            ("keep", "- **非定制类产品**：……"),
        ],
        tmpl,
    )
    p = find_exact(doc, "    - **定制类产品**：……")
    replace_para(
        p,
        [
            ("ins", "    （仅路径B）"),
            ("keep", "- **定制类产品**：……"),
        ],
        tmpl,
    )

    p = find_exact(
        doc,
        "报告章节标题**只能**用四个固定名：`【核心风险诊断】` `【合规方案】` `【行动建议】` `【注意事项】`。",
    )
    insert_after(
        p,
        [[
            (
                "ins",
                "禁止写成「第一部分／第二部分／第四部分／第五部分」或「核心风险提示」「合规方案建议」等变体；不要讨论章节编号是否缺第三部分。",
            )
        ]],
        tmpl,
    )

    p = find_exact(
        doc,
        "5. 非定制／定制对照写成**两条并列卡片**（不要嵌套 a./b.）：",
    )
    replace_para(
        p,
        [
            ("keep", "5. 非定制／定制对照写成**两条并列卡片**（不要嵌套 a./b.）"),
            ("ins", "——仅路径B需要拆非定制/定制时使用；路径A/C/D禁止套用这两条对照卡片"),
            ("keep", "："),
        ],
        tmpl,
    )

    p = find_exact(doc, "未达档位不必硬写；达到则【合规方案】或【注意事项】至少一条覆盖。")
    insert_after(
        p,
        [
            [("ins", "---")],
            [("ins", "## 内置合规架构速查（检索失败时的兜底；禁止对用户输出本节标题与路径编号）")],
            [(
                "ins",
                "选型时结合用户平台/发货/出口/发票/产品类别/销售额；必读库有更细平台分支时优先采信检索结果，检索空缺或冲突时用本速查给出可执行路径，禁止空泛作答。",
            )],
            [("ins", "1. **「0110出口+香港公司」**")],
            [("ins", "链路：国内供应商 → 进出口公司 → 香港公司 → 店铺公司 → 境外消费者。")],
            [(
                "ins",
                "进出口公司0110报关出口给香港公司（香港公司为境外买家）；供应商开票给进出口公司，专票可退税、普票多走免税。关联公司注意转让定价，勿把大部分利润留在无实质经营的香港公司。",
            )],
            [(
                "ins",
                "**主体一致性**：架构搭好后须核对香港公司是否=平台店铺主体；若否（常见店群），且出口非「由平台安排出口／全托管国内仓」等，须补「香港公司向出口主体采购后再销售给店铺公司」链路。全托管等平台安排出口场景不适用该采购再销售链路。",
            )],
            [("ins", "2. **「1039出口+香港公司」**")],
            [("ins", "链路：国内供应商 → 个体户 → 香港公司 → 店铺公司 → 境外消费者。")],
            [(
                "ins",
                "个体户在市场采购区（常建议东莞/义乌）1039出口给香港公司；供应商可不开发票；个体户可享无票免征增值税 + 个税核定，但连续12个月销售额一般不超过500万，量大需多个体户。同样注意关联交易/避税认定。",
            )],
            [(
                "ins",
                "1039须特别注意：①敏感产品、危险及化工产品、0退税产品、液体、动植物食品类不能走1039；②商标或专利须提前核海关知识产权备案，已备案须获授权才可出口；③含木产品需提前商检，原木（实木）包装要提前熏蒸。",
            )],
            [("ins", "3. **「1210出口备货至保税区」**")],
            [("ins", "链路：国内供应商 → 进出口公司 → 备货保税区（货主香港公司）→ 店铺公司 → 境外消费者。")],
            [(
                "ins",
                "先0110进保税区给香港公司，再按平台订单1210一件代发离境；专票可争取进区即退税（部分税局仍要离境清单）。仅适合非定制类，定制易滞销。",
            )],
            [("ins", "4. **「1210保税区一日游或9610跨境电商零售出口」**")],
            [("ins", "链路：国内供应商 → 进出口公司（常即店铺公司）→ 保税区/监管场站（货主进出口公司）→ 境外消费者。")],
            [(
                "ins",
                "订单驱动1210/9610报关；专票可退税、普票多免税。需商检商品通常不适合（逐包裹商检成本高）。",
            )],
        ],
        tmpl,
    )

    p = find_exact(
        doc,
        "有票 / 无票货源拆分：部分专票+部分无票时，有票走退税架构、无票单独评估「1039出口+香港公司」（额度、禁限类、知识产权备案），勿混成一句含糊建议。",
    )
    insert_after(
        p,
        [[
            (
                "ins",
                "1039禁限（写无票/1039时必须核）：①敏感产品、危险及化工产品、0退税产品、液体、动植物食品类不能走1039；②商标或专利须提前核海关知识产权备案，已备案须获授权才可出口；③含木产品需提前商检，原木（实木）包装要提前熏蒸。",
            )
        ]],
        tmpl,
    )

    p = find_startswith(doc, "解决方案的核心不是如何满足正式报关出口或者退税条件")
    replace_para(
        p,
        [
            (
                "keep",
                "解决方案的核心不是如何满足正式报关出口或者退税条件，而是需要注意合理安排业务主体，",
            ),
            (
                "del",
                "例如保持小规模纳税人身份，视同内销增值税1%，税负可控；如果连续12个月超过500万，需按一般纳税人按13%缴纳增值税，税负太高，建议考虑香港主体",
            ),
            (
                "ins",
                "仅当年销售额＜500万时，可写「保持小规模纳税人身份，视同内销增值税1%，税负相对可控」；年销售额≥500万时，禁止写保持小规模，须按一般纳税人视同内销征税，并评估引入香港等境外主体作为销售主体。",
            ),
        ],
        tmpl,
    )

    save(doc, "出报告硬约束与路径要点-修订对照.docx")


def build_sample():
    doc = Document(SRC / "23-____-__D-_________655a.docx")
    tmpl = body_template_run(doc)
    add_legend(doc, tmpl)

    p = find_exact(doc, "•\t正视视同内销增值税")
    replace_para(
        p,
        [
            ("del", "•\t正视视同内销增值税"),
            (
                "ins",
                "- **正视视同内销税负**：测算视同内销增值税税负，评估利润空间，必要时需要将增值税税负考虑到产品定价，一定程度上降低合规税负带来的经营压力。",
            ),
        ],
        tmpl,
    )
    p = find_startswith(doc, "测算视同内销增值税税负，评估利润空间")
    replace_para(
        p,
        [("del", p.text)],
        tmpl,
    )

    p = find_exact(doc, "•\t合理安排业务主体")
    replace_para(
        p,
        [
            ("del", "•\t合理安排业务主体"),
            (
                "ins",
                "- **合理安排业务主体**：解决方案的核心不是如何满足正式报关出口或者退税条件，而是需要注意合理安排业务主体。目前销售规模已经超过500万，已难维持小规模，需按一般纳税人视同内销征税，可评估引入香港等境外主体，重新安排业务流程。",
            ),
        ],
        tmpl,
    )
    p = find_startswith(doc, "解决方案的核心不是如何满足正式报关出口或者退税条件")
    replace_para(
        p,
        [("del", p.text)],
        tmpl,
    )

    p = find_exact(doc, "•\t必要时重构业务模式")
    replace_para(
        p,
        [
            ("del", "•\t必要时重构业务模式"),
            (
                "ins",
                "- **必要时重构业务模式**：部分行业可以通过业务模式重构降低合规税负，例如通过原材料保税进口，加工生产后复出口的方式，可以合理避免增值税视同内销的风险，降低合规税负。",
            ),
        ],
        tmpl,
    )
    p = find_startswith(doc, "部分行业可以通过业务模式重构降低合规税负")
    replace_para(
        p,
        [("del", p.text)],
        tmpl,
    )

    save(doc, "23-方案样本-路径D-修订对照.docx")


if __name__ == "__main__":
    build_prompt()
    build_hard()
    build_sample()
    print("done")
